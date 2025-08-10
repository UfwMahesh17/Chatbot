
import os
import io
import re
import time
import random
import logging
import unicodedata
import string
from collections import deque
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename

from langchain_community.vectorstores import Chroma
from langchain_cohere import CohereEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter

import PyPDF2
from docx import Document as DocxDocument
import cohere

# ---------- Config ----------
COHERE_API_KEY = "IONaj3UR5GCuv2EBIVrVvWEe8W6UEllHmX5MzbKz"


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploaded_files")
PERSIST_DIR = os.path.join(BASE_DIR, "chroma_db")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Contact text (single source of truth)
CONTACT_LINE = "support@agent42labs.com or +91 7027119799"
CONTACT_SENTENCE = f"Contact us at {CONTACT_LINE}."
CONTACT_MORE_INFO = f"For more info, contact us at {CONTACT_LINE}."

# Retrieval / Re-ranking
RERANK_MODEL = "rerank-english-v3.0"
INITIAL_K = 40            # initial candidates
FINAL_DOCS = 6            # pass best N to LLM
RERANK_THRESHOLD = 0.35   # gate weak matches (raise to 0.50–0.60 to be stricter)

# Chunking
splitter = RecursiveCharacterTextSplitter(
    chunk_size=800,
    chunk_overlap=100,
    separators=["\n\n", "\n", ". ", "! ", "? ", " "]
)

ALLOWED_EXTENSIONS = {"pdf", "txt", "docx"}

# ---------- App / Clients ----------
app = Flask(__name__)
CORS(app)

embeddings = CohereEmbeddings(model="embed-english-v3.0", cohere_api_key=COHERE_API_KEY)
vectorstore = Chroma(
    collection_name="langchain",
    persist_directory=PERSIST_DIR,
    embedding_function=embeddings
)
co = cohere.Client(COHERE_API_KEY)

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

# ---------- Professional fallback rotation ----------
FALLBACK_MESSAGES = [
    "I couldn’t find that in the available materials.",
    "This information isn’t present in the content I have access to.",
    "I don’t have enough detail in the current documents to answer that.",
    "That topic doesn’t appear in our indexed content.",
    "I’m not seeing a direct match for that in the context.",
    "The current documents don’t cover that request.",
    "I wasn’t able to locate a source for that in the materials.",
    "I can’t confirm that from the context I have.",
    "There isn’t sufficient information in the indexed content to answer that confidently.",
    "I couldn’t verify that in the available sources.",
    "I don’t have a documented answer for that in the knowledge base.",
    "That doesn’t seem to be covered in the material I’m using.",
    "The context I have doesn’t include an answer to that.",
    "I didn’t find supporting details for that in the sources.",
    "I’m not able to provide a context‑backed answer to that.",
    "It looks like this isn’t documented in the materials I have.",
    "I wasn’t able to find a relevant reference for that.",
    "I don’t have the specifics for that in the current set of documents.",
    "The information you’re looking for isn’t available in the indexed content.",
    "I couldn’t locate a reliable source for that in the context."
]
REFINE_HINTS = [
    "If you can share a bit more detail (e.g., specific service, product, or page), I’ll look again.",
    "Please add more context—such as the area or page you’re referring to—and I’ll recheck.",
    "Point me to a page title or file name if you can, and I’ll search that directly.",
    "If you specify the topic or section, I’ll try a more targeted search.",
    "Share any keywords or the exact phrase you saw, and I’ll look it up.",
    "Let me know the timeframe or team (e.g., Services, Industries), and I’ll refine the search."
]
_recent_base = deque(maxlen=6)
_recent_hint = deque(maxlen=3)

def _pick_non_recent(options, recent, attempts=10):
    for _ in range(attempts):
        choice = random.choice(options)
        if choice not in recent or len(set(options)) <= len(recent):
            recent.append(choice)
            return choice
    recent.clear()
    choice = random.choice(options)
    recent.append(choice)
    return choice

def get_fallback_response(fail_count=0):
    base = _pick_non_recent(FALLBACK_MESSAGES, _recent_base)
    hint = _pick_non_recent(REFINE_HINTS, _recent_hint)
    msg = f"{base}\n\n{hint}"
    if fail_count >= 2:
        msg = f"{msg}\n\n{CONTACT_SENTENCE}"
    return msg

# ---------- Greetings / thanks / goodbye ----------
GREETING_PHRASES = {
    "hi", "hello", "hey", "greetings", "hi there", "hello there", "hey there",
    "good morning", "good afternoon", "good evening", "morning", "afternoon", "evening"
}
THANKS_PHRASES = {"thanks", "thank you", "thankyou", "thx", "ty", "much appreciated", "appreciate it"}
GOODBYE_PHRASES = {"bye", "goodbye", "see you", "see ya", "talk later", "talk to you later"}

GREETING_REPLY = (
    "Hello! I’m the Agent42 Labs assistant. I can help with information about our services, industries, "
    "solutions, and company details, or connect you with the team. How can I assist you today?\n\n"
    "You can ask about:\n- Services we offer\n- Industries we work with\n- Capabilities or case studies\n- How to contact our team"
)
THANKS_REPLY = "You’re welcome—happy to help. If there’s anything else you need, let me know."
GOODBYE_REPLY = f"Thanks for chatting. If you need anything later, {CONTACT_SENTENCE}"

def is_greeting(norm_q: str) -> bool:
    q = norm_q.strip()
    return q in GREETING_PHRASES or any(q.startswith(p) for p in GREETING_PHRASES)

def is_thanks(norm_q: str) -> bool:
    return any(phrase in norm_q for phrase in THANKS_PHRASES)

def is_goodbye(norm_q: str) -> bool:
    return any(phrase in norm_q for phrase in GOODBYE_PHRASES)

# ---------- Utilities ----------
def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def uniquify(path: str) -> str:
    if not os.path.exists(path):
        return path
    base, ext = os.path.splitext(path)
    i = 1
    while os.path.exists(f"{base}_{i}{ext}"):
        i += 1
    return f"{base}_{i}{ext}"

def normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\u00A0", " ")
    text = re.sub(r"(?m)^\s*\.\s*$", "", text)              # lone dots
    text = re.sub(r"[ \t]+", " ", text)                     # collapse spaces
    text = re.sub(r"\n{3,}", "\n\n", text)                  # collapse blank lines
    return text.strip()

def extract_text_from_file(filepath):
    ext = filepath.rsplit(".", 1)[-1].lower()
    try:
        if ext == "txt":
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return normalize_text(f.read())
        elif ext == "pdf":
            text = []
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    t = page.extract_text() or ""
                    text.append(t)
            return normalize_text("\n".join(text))
        elif ext == "docx":
            doc = DocxDocument(filepath)
            return normalize_text("\n".join([p.text for p in doc.paragraphs]))
    except Exception as e:
        logging.error(f"Error extracting text from {filepath}: {e}")
    return ""

def sanitize_contact_text(text: str) -> str:
    text = re.sub(r"contact us at\s+you can reach us at", "contact us at", text, flags=re.I)
    if text.count(CONTACT_LINE) > 1:
        first, rest = text.split(CONTACT_LINE, 1)
        rest = rest.replace(CONTACT_LINE, "")
        text = first + CONTACT_LINE + rest
    return text

def normalize_question(q):
    return q.lower().translate(str.maketrans("", "", string.punctuation)).strip()

# ---------- Pricing guardrail (return contact info immediately) ----------
PRICING_KEYWORDS = [
    "price", "pricing", "cost", "quote", "quotation", "estimate", "budget",
    "rate", "rates", "fee", "fees", "charge", "charges",
    "how much", "per month", "per user", "per seat", "subscription", "plan", "plans", "rate card"
]
CURRENCY_RE = re.compile(r"(?:[$€£]|₹|rs\.?)\s?\d[\d,]*(?:\.\d+)?|\d[\d,]*(?:\.\d+)?\s?(?:usd|inr|eur|gbp|rs\.?|rupees)", re.I)

def is_pricing_intent(raw_q: str, norm_q: str) -> bool:
    if any(k in norm_q for k in PRICING_KEYWORDS):
        return True
    return bool(CURRENCY_RE.search(raw_q))

# ---------- Routes ----------
@app.route("/upload", methods=["POST"])
def upload_file():
    try:
        files = request.files.getlist("files")
        if not files:
            return jsonify({"error": "No files uploaded"}), 400

        total_chunks = 0
        for file in files:
            original_name = file.filename or ""
            if not original_name:
                continue
            if not allowed_file(original_name):
                logging.warning(f"Skipped unsupported file: {original_name}")
                continue

            safe_name = secure_filename(original_name).replace(" ", "_")
            if not safe_name:
                safe_name = f"file_{int(time.time()*1000)}.txt"

            filepath = uniquify(os.path.join(UPLOAD_FOLDER, safe_name))
            file.save(filepath)

            text = extract_text_from_file(filepath)
            if not text:
                continue

            chunks = splitter.split_text(text)
            if not chunks:
                continue

            # Upsert for this file: delete old vectors for this source, then add fresh
            try:
                vectorstore._collection.delete(where={"source": safe_name})  # type: ignore
            except Exception as e:
                logging.info(f"No prior vectors to delete for {safe_name}: {e}")

            metas = [{"source": safe_name, "original_filename": original_name}] * len(chunks)
            vectorstore.add_texts(chunks, metadatas=metas)
            total_chunks += len(chunks)

        try:
            vectorstore.persist()
        except Exception as e:
            logging.warning(f"Persist warning: {e}")

        return jsonify({"message": f"Files uploaded and indexed successfully. Chunks added: {total_chunks}"})
    except Exception as e:
        logging.error(f"Upload error: {e}")
        return jsonify({"error": "Internal server error"}), 500

@app.route("/chat", methods=["POST"])
def chat():
    data = request.json or {}
    question = (data.get("question") or "").strip()
    history = data.get("history", [])
    fail_count = int(data.get("fail_count", 0))

    if not question:
        return jsonify({"error": "No question provided"}), 400

    norm_q = normalize_question(question)

    # 0) Pricing guardrail
    if is_pricing_intent(question, norm_q):
        return jsonify({"answer": CONTACT_SENTENCE, "fail_count": fail_count})

    # 0.1) Greetings / thanks / goodbye
    if is_greeting(norm_q):
        return jsonify({"answer": GREETING_REPLY, "fail_count": fail_count})
    if is_thanks(norm_q):
        return jsonify({"answer": THANKS_REPLY, "fail_count": fail_count})
    if is_goodbye(norm_q):
        return jsonify({"answer": GOODBYE_REPLY, "fail_count": fail_count})

    # 1) Contact intent (general)
    contact_keywords = [
        "contact", "contact info", "contact information", "email", "phone", "phone number",
        "address", "support", "customer service", "how do i contact", "how can i contact",
        "get in touch", "reach you", "reach agent42", "contact agent42"
    ]
    if any(k in norm_q for k in contact_keywords):
        return jsonify({"answer": CONTACT_SENTENCE, "fail_count": fail_count})

    # 2) Retrieve candidates (MMR for diversity)
    try:
        raw_docs = vectorstore.max_marginal_relevance_search(
            question, k=INITIAL_K, fetch_k=INITIAL_K * 2, lambda_mult=0.2
        )
    except Exception:
        raw_docs = vectorstore.similarity_search(question, k=INITIAL_K)

    # 3) Rerank with Cohere
    selected_texts = []
    top_score = 0.0
    try:
        docs_text = [d.page_content for d in raw_docs]
        rr = co.rerank(
            model=RERANK_MODEL,
            query=question,
            documents=docs_text,
            top_n=min(FINAL_DOCS, len(docs_text)) if docs_text else 0
        )
        selected = []
        for r in rr.results or []:
            d = raw_docs[r.index]
            header = []
            sec = d.metadata.get("section")
            it = d.metadata.get("item_title")
            if sec: header.append(sec)
            if it: header.append(it)
            prefix = " — ".join(header)
            selected.append(f"{prefix}\n{d.page_content}" if prefix else d.page_content)
        selected_texts = selected
        if rr.results:
            top_score = rr.results[0].relevance_score or 0.0
    except Exception as e:
        logging.warning(f"Rerank failed, falling back to top-k: {e}")
        selected_texts = [d.page_content for d in raw_docs[:FINAL_DOCS]]

    # 4) Confidence gate
    if not selected_texts or top_score < RERANK_THRESHOLD:
        answer = get_fallback_response(fail_count)
        fail_count += 1
        answer = sanitize_contact_text(answer)
        return jsonify({"answer": answer, "fail_count": fail_count})

    context = "\n\n---\n\n".join(selected_texts)

    prompt = f"""
You are a helpful assistant for Agent42 Labs. Answer ONLY using the provided context below.
If the context contains a list, enumerate all items in the list.
Summarize and organize your answer clearly, avoid repetition, and be professional and friendly.
If the answer is not in the context, do NOT make up information. If you cannot answer, say so and offer to connect the user to support.

Context:
{context}

Question:
{question}

Answer:
""".strip()

    try:
        response = co.generate(
            model="command-r-plus",
            prompt=prompt,
            max_tokens=500,
            temperature=0.1
        )
        answer = (response.generations[0].text or "").strip()
    except Exception as e:
        logging.error(f"Cohere generate error: {e}")
        answer = ""

    # 5) Fallback if weak/empty
    if (not answer) or ("i don't know" in answer.lower()) or ("no information" in answer.lower()) or answer.strip().endswith(":"):
        answer = get_fallback_response(fail_count)
        fail_count += 1
    else:
        fail_count = 0

    answer = sanitize_contact_text(answer)
    logging.info(f"Q: {question} | A: {answer}")
    return jsonify({"answer": answer, "fail_count": fail_count})

@app.route("/export", methods=["POST"])
def export():
    try:
        data = request.json or {}
        export_type = data.get("type", "txt")
        content = data.get("content", "")
        filename = data.get("filename", "chat_export")

        if export_type == "txt":
            return send_file(io.BytesIO(content.encode("utf-8")),
                             mimetype="text/plain",
                             as_attachment=True,
                             download_name=f"{filename}.txt")
        elif export_type == "json":
            return send_file(io.BytesIO(content.encode("utf-8")),
                             mimetype="application/json",
                             as_attachment=True,
                             download_name=f"{filename}.json")
        elif export_type == "pdf":
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=12)
            for line in content.split("\n"):
                pdf.multi_cell(0, 8, line)
            pdf_bytes = pdf.output(dest="S").encode("latin-1")
            return send_file(io.BytesIO(pdf_bytes),
                             mimetype="application/pdf",
                             as_attachment=True,
                             download_name=f"{filename}.pdf")
        elif export_type == "docx":
            from docx import Document
            doc = Document()
            for line in content.split("\n"):
                doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return send_file(buf,
                             mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             as_attachment=True,
                             download_name=f"{filename}.docx")
        else:
            return jsonify({"error": "Invalid export type"}), 400
    except Exception as e:
        logging.error(f"Export error: {e}")
        return jsonify({"error": "Export failed"}), 500

if __name__ == "__main__":
    app.run(port=5000, debug=False)